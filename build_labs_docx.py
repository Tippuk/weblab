# -*- coding: utf-8 -*-
"""Скрипт для генерации сводки по лабораторным работам в формате DOCX."""
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Установите python-docx: pip install python-docx")
    sys.exit(1)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_table_from_rows(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
        t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            t.rows[i + 1].cells[j].text = str(cell)
    doc.add_paragraph()

def main():
    doc = Document()
    doc.add_heading("Сводка по лабораторным работам — проект Новостник", 0)

    # ========== ЛАБА 2 ==========
    add_heading(doc, "Лабораторная работа 2 — Базовое приложение, шаблоны, статика, формы", 1)
    add_paragraph(doc, "Темы: базовое приложение, базовый шаблон, статические файлы, форма обратной связи, динамические маршруты, меню, список статей, значок «Новое!».")

    add_heading(doc, "Где и как реализовано (Лаба 2)", 2)
    rows_laba2 = [
        ("A. Базовое приложение", "app.py", "Маршрут / — index() (стр. 115–124): главная с пагинацией. Текст «Добро пожаловать в Новостник!» в templates/index.html."),
        ("Обработчики /about, /contact", "app.py", "about() (126–128), contact() (131–133); рендер about.html, contact.html."),
        ("B. Базовый шаблон", "templates/base.html", "Общий каркас: head, стили, header с меню, main с блоком content, footer. Наследование через extends."),
        ("C. Статические файлы", "static/css/style.css, base.html", "Подключение: url_for('static', filename='css/style.css'). Логотип — текстовая ссылка «News Hub» в header."),
        ("D. Форма обратной связи", "app.py, templates/feedback.html", "Маршрут /feedback (136–156): GET — форма, POST — валидация (имя, email, сообщение). После отправки — вывод данных."),
        ("E. Динамический маршрут", "app.py", "Маршрут /news/<int:id> → article(id) (198–202): вывод статьи по id или 404. Шаблон article.html."),
        ("F. Навигационное меню", "templates/base.html", "main-nav: ссылки на главную, статьи, создать статью, вход/выход, о проекте, контакты, обратная связь."),
        ("G. Список статей на главной", "app.py, index.html", "В index() передаётся articles, pagination. В шаблоне цикл по статьям — заголовок и дата."),
        ("H. Значок «Новое!»", "app.py, index.html, article.html", "template_global is_today(d) в app.py (56–62). В шаблонах: if is_today(article.created_date) — badge «Новое!»."),
    ]
    add_table_from_rows(doc, ["Тема", "Файл/место", "Как реализовано"], rows_laba2)

    # ========== ЛАБА 3 ==========
    add_heading(doc, "Лабораторная работа 3 — Модели, CRUD, фильтрация, комментарии, аутентификация", 1)
    add_paragraph(doc, "Темы: модели User и Article, связь один-ко-многим, CRUD статей, страница /articles и категории, модель Comment и форма комментариев, регистрация и вход, защита создания/редактирования статей.")

    add_heading(doc, "Где и как реализовано (Лаба 3)", 2)
    rows_laba3 = [
        ("A. Модели User и Article", "models.py", "User: id, name, email, hashed_password, created_date. Article: id, title, text, created_date, user_id, category. Связь: User.articles, Article.author (relationship, ForeignKey)."),
        ("B. CRUD статей", "app.py, шаблоны", "Создание: /create-article (221–247), create_article.html. Редактирование: /edit-article/<id> (250–263). Удаление: /delete-article/<id> (266–272). Везде @login_required."),
        ("C. /articles и категории", "app.py", "/articles — articles_list() с поиском q. /articles/<category> — articles_by_category(), фильтр по category, при неверной категории 404."),
        ("D. Комментарии", "models.py, app.py, article.html", "Модель Comment (id, text, date, article_id, author_name). Форма в article.html, POST /news/<id>/comment → add_comment(). Вывод comments в article()."),
        ("E. Аутентификация", "app.py, models.py", "Регистрация /register, вход /login (User.check_password, login_user). Создание/редактирование/удаление статей только для авторизованных (@login_required)."),
    ]
    add_table_from_rows(doc, ["Тема", "Файл/место", "Как реализовано"], rows_laba3)

    # ========== ЛАБА 4 ==========
    add_heading(doc, "Лабораторная работа 4 — API: эндпоинты, CRUD, фильтрация, комментарии", 1)
    add_paragraph(doc, "Темы: GET /api/articles и /api/articles/<id>, POST/PUT/DELETE статей с валидацией, фильтр по категории и сортировка по дате, CRUD для Comment (в лабе — /api/comment; в проекте — комментарии к статье).")

    add_heading(doc, "A. Базовые эндпоинты (Лаба 4)", 2)
    rows_4a = [
        ("GET /api/articles", "app.py, 432–455", "api_articles_list(): пагинация page, per_page, поиск q, фильтр category. Ответ: articles, page, per_page, total, pages."),
        ("GET /api/articles/<id>", "app.py, 457–460", "api_article_get(id): get_or_404, ответ article_to_dict(a)."),
    ]
    add_table_from_rows(doc, ["Эндпоинт", "Где", "Как"], rows_4a)

    add_heading(doc, "B. CRUD через API (Лаба 4)", 2)
    rows_4b = [
        ("POST /api/articles", "app.py, 496–511", "api_article_create(): проверка JWT. Валидация: title и text обязательны, category из ARTICLE_CATEGORIES. sanitize_html."),
        ("PUT /api/articles/<id>", "app.py, 515–528", "api_article_update(): обновление title, text, category (только из списка категорий)."),
        ("DELETE /api/articles/<id>", "app.py, 532–539", "api_article_delete(): проверка авторизации, delete статьи, ответ 204."),
    ]
    add_table_from_rows(doc, ["Эндпоинт", "Где", "Как"], rows_4b)

    add_heading(doc, "C. Фильтрация и сортировка (Лаба 4)", 2)
    rows_4c = [
        ("Фильтр по категории", "api_articles_list()", "Query-параметр: GET /api/articles?category=tech. Проверка category in ARTICLE_CATEGORIES."),
        ("Сортировка по дате", "api_articles_list()", "Фиксированная: order_by(Article.created_date.desc()). Отдельного параметра sort нет."),
    ]
    add_table_from_rows(doc, ["Требование", "Где", "Как"], rows_4c)

    add_heading(doc, "D. CRUD для Comment (Лаба 4)", 2)
    add_paragraph(doc, "В лабе: /api/comment и /api/comment/<id>. В проекте комментарии привязаны к статье: /api/articles/<id>/comments.")
    rows_4d = [
        ("GET все комментарии", "Нет отдельного", "Есть GET /api/articles/<id>/comments — комментарии одной статьи (api_article_comments, 473–477)."),
        ("GET комментарий по ID", "Нет", "—"),
        ("POST комментарий", "app.py, 480–492", "POST /api/articles/<id>/comments — api_article_add_comment. Валидация: text и author_name обязательны."),
        ("PUT комментарий", "Нет", "—"),
        ("DELETE комментарий", "Нет", "—"),
    ]
    add_table_from_rows(doc, ["Требование", "В проекте", "Примечание"], rows_4d)

    # ========== ЛАБА 5 ==========
    add_heading(doc, "Лабораторная работа 5 — JWT", 1)
    add_paragraph(doc, "Темы: эндпоинты access/refresh, интеграция JWT, middleware проверки токенов, обновление токенов.")

    add_heading(doc, "Где и как реализовано (Лаба 5)", 2)
    rows_laba5 = [
        ("Эндпоинты access и refresh", "app.py", "POST /api/auth/login (350–366): create_access_token, create_refresh_token. POST /api/auth/refresh (369–375): @jwt_required(refresh=True), новый access_token."),
        ("Интеграция JWT", "app.py", "flask_jwt_extended, JWT_SECRET_KEY, JWT_ACCESS_TOKEN_EXPIRES (минуты), JWT_REFRESH_TOKEN_EXPIRES (дни) — стр. 27–29. Регистрация API тоже выдаёт токены (386–406)."),
        ("Проверка JWT (middleware)", "app.py", "get_api_user() (311–322): verify_jwt_in_request(optional=True), get_jwt_identity(). Используется в api_article_create/update/delete. /api/auth/me — jwt_required(optional=True)."),
        ("Обновление токенов", "Бэкенд + фронт", "Бэкенд: /api/auth/refresh возвращает новый access. Фронт: в api.js при 401 и наличии refresh — запрос на /auth/refresh, setStoredTokens, повтор запроса."),
    ]
    add_table_from_rows(doc, ["Тема", "Где", "Как"], rows_laba5)

    add_paragraph(doc, "Срок действия токенов (app.py, 28–29): JWT_ACCESS_TOKEN_EXPIRES — по умолчанию 15 минут (переменная JWT_ACCESS_TOKEN_EXPIRES в минутах); JWT_REFRESH_TOKEN_EXPIRES — 7 дней (JWT_REFRESH_TOKEN_EXPIRES в днях).")

    # ========== ЛАБА 6 ==========
    add_heading(doc, "Лабораторная работа 6 — Фронтенд, работа с API", 1)
    add_paragraph(doc, "Темы: новостной сайт на фронте, авторизация по JWT, без UI-библиотек, Git, линтер/форматтер, при необходимости — доработка API.")

    add_heading(doc, "Где и как реализовано (Лаба 6)", 2)
    rows_laba6 = [
        ("Новостной сайт", "frontend/src", "React (Vite): Home.jsx, Article.jsx, CreateArticle.jsx, EditArticle.jsx. Роутинг в App.jsx."),
        ("Авторизация по JWT", "frontend/src/api.js, AuthContext.jsx", "Токены в localStorage. Authorization: Bearer <access>. При 401 — обновление через /auth/refresh. AuthProvider, useAuth, ProtectedRoute в App.jsx. Login.jsx, Register.jsx."),
        ("Без UI-библиотек", "frontend", "Стили в assets/style.css, без Material UI и т.п."),
        ("Git", "Корень проекта", "Репозиторий настроен."),
        ("Линтер/форматтер", "frontend", "Конфиг .eslintrc.cjs."),
        ("Эндпоинты API", "app.py", "auth/login, register, refresh, me; articles GET/POST, articles/<id> GET/PUT/DELETE; articles/<id>/comments GET/POST; /api/feedback POST."),
    ]
    add_table_from_rows(doc, ["Тема", "Где", "Как"], rows_laba6)

    doc.add_paragraph()
    add_paragraph(doc, "Итог: статьи и CRUD по API (Лаба 4) реализованы; фильтр по категории — через query ?category=...; сортировка по дате — фиксированная. Комментарии в проекте — только в контексте статьи (GET/POST по article id); глобальные /api/comment и PUT/DELETE комментария по ID в проекте отсутствуют.", bold=True)

    out_path = "Сводка_по_лабам.docx"
    doc.save(out_path)
    print(f"Файл сохранён: {out_path}")

if __name__ == "__main__":
    main()
